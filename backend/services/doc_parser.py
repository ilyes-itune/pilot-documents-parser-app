from typing import Dict, Any, List
import io
try:
    from docx import Document
except ImportError:
    Document = None

class DocParser:
    def parse(self, file_content: bytes) -> Dict[str, Any]:
        if Document is None:
            raise ImportError("python-docx required. Install: pip install python-docx")
        
        try:
            doc = Document(io.BytesIO(file_content))
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            tables = self._extract_tables(doc)
            metadata = self._extract_metadata(doc)
            
            return {
                'text': text,
                'tables': tables,
                'author': metadata.get('author'),
                'title': metadata.get('title'),
                'subject': metadata.get('subject'),
                'created': metadata.get('created'),
                'modified': metadata.get('modified')
            }
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")

    @staticmethod
    def _extract_tables(doc) -> List[Dict[str, Any]]:
        tables = []
        for table in doc.tables:
            headers = []
            rows = []
            
            if len(table.rows) > 0:
                for cell in table.rows[0].cells:
                    headers.append(cell.text)
            
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                rows.append(row_data)
            
            if headers or rows:
                tables.append({'headers': headers, 'rows': rows})
        
        return tables

    @staticmethod
    def _extract_metadata(doc) -> Dict[str, Any]:
        metadata = {}
        if hasattr(doc, 'core_properties'):
            props = doc.core_properties
            metadata['author'] = props.author
            metadata['title'] = props.title
            metadata['subject'] = props.subject
            metadata['created'] = str(props.created) if props.created else None
            metadata['modified'] = str(props.modified) if props.modified else None
        return metadata